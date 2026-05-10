import fitz
from docx import Document
import openpyxl


def parse_file(filepath):
    text = ""

    if filepath.endswith('.pdf'):
        text = parse_pdf(filepath)
    elif filepath.endswith('.docx'):
        text = parse_docx(filepath)
    elif filepath.endswith('.xlsx') or filepath.endswith('.xls'):
        text = parse_excel(filepath)
    elif filepath.endswith('.txt') or filepath.endswith('.csv'):
        with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
            text = f.read()

    return text  # No limit — full document


def parse_pdf(filepath):
    text = ""
    doc = fitz.open(filepath)
    for page_num, page in enumerate(doc):
        page_text = page.get_text()
        if page_text.strip():
            text += f"\n--- Page {page_num + 1} ---\n"
            text += page_text
    return text


def parse_docx(filepath):
    doc = Document(filepath)
    text = ""

    for para in doc.paragraphs:
        if para.text.strip():
            if para.style.name.startswith('Heading'):
                text += f"\n=== {para.text.strip()} ===\n"
            else:
                text += para.text.strip() + "\n"

    # Also read tables inside DOCX
    for table in doc.tables:
        text += "\n[TABLE]\n"
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                text += row_text + "\n"
        text += "[END TABLE]\n"

    return text


def parse_excel(filepath):
    wb = openpyxl.load_workbook(filepath)
    text = ""
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        text += f"\n=== Sheet: {sheet_name} ===\n"
        for row in ws.iter_rows(values_only=True):
            row_text = " | ".join(
                [str(c) for c in row if c is not None]
            )
            if row_text.strip():
                text += row_text + "\n"
    return text


def split_into_chunks(text, chunk_size=8000, overlap=500):
    """
    Splits large document into overlapping chunks.
    overlap keeps context between chunks so nothing is missed.
    """
    if len(text) <= chunk_size:
        return [text]

    chunks = []
    start = 0

    while start < len(text):
        end = start + chunk_size

        if end < len(text):
            # Cut at paragraph boundary instead of mid-sentence
            newline_pos = text.rfind('\n', start, end)
            if newline_pos > start + (chunk_size // 2):
                end = newline_pos

        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)

        start = end - overlap
        if start >= len(text):
            break

    return chunks


def get_document_stats(text):
    chunks = split_into_chunks(text)
    return {
        "total_characters": len(text),
        "total_words": len(text.split()),
        "total_chunks": len(chunks),
        "estimated_pages": round(len(text) / 3000)
    }